#!/usr/bin/env python3
"""
Genera el informe final en formato DOCX a partir de los resúmenes
generados por cada script de algoritmo y el contenido del README.

Embebe todas las imágenes con pies de figura y tablas con cabeceras,
siguiendo el formato de reporte científico.

Uso:
  python scripts/generar_docx.py
  python scripts/generar_docx.py --output informe_final.docx
"""
from __future__ import annotations

import argparse
import os
import re

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


def add_cover(doc: Document) -> None:
    """Portada del informe."""
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Universidad del Valle de Guatemala")
    run.font.size = Pt(16)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Minería de Datos")
    run.font.size = Pt(14)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Tarea 2: Otros Algoritmos de Aprendizaje No Supervisado")
    run.font.size = Pt(18)
    run.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Integrantes del grupo:")
    run.font.size = Pt(12)
    run.italic = True

    names = ["[Nombre 1]", "[Nombre 2]", "[Nombre 3]"]
    for name in names:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(name).font.size = Pt(12)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Febrero 2026")
    run.font.size = Pt(12)

    doc.add_page_break()


def parse_md_table(lines: list[str]) -> list[list[str]]:
    """Parsea líneas de tabla markdown a lista de filas."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        cells = [c.strip() for c in line.split("|")]
        cells = [c for c in cells if c]
        if all(set(c) <= {"-", ":"} for c in cells):
            continue
        rows.append(cells)
    return rows


def add_table_to_doc(doc: Document, rows: list[list[str]],
                     caption: str = "") -> None:
    """Agrega una tabla con formato al documento."""
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.bold = True
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if not rows:
        return

    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < n_cols:
                cell = table.cell(i, j)
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    paragraph.style.font.size = Pt(9)
                    if i == 0:
                        for run in paragraph.runs:
                            run.bold = True

    doc.add_paragraph()


def add_image_to_doc(doc: Document, img_path: str, caption: str,
                     width: float = 5.5) -> None:
    """Agrega una imagen centrada con pie de figura."""
    if not os.path.isfile(img_path):
        p = doc.add_paragraph(f"[Imagen no encontrada: {img_path}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Inches(width))

    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        run.font.size = Pt(9)
        run.italic = True
    doc.add_paragraph()


def process_md_content(doc: Document, md_text: str, img_dir: str,
                       heading_offset: int = 0) -> None:
    """Procesa contenido markdown y lo agrega al documento DOCX.

    Maneja: headings, bold, listas, tablas, imágenes y párrafos.
    heading_offset ajusta el nivel de los headings (0 = sin cambio).
    """
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Líneas vacías
        if not stripped:
            i += 1
            continue

        # Imagen markdown: ![alt](file.png)
        img_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if img_match:
            img_file = img_match.group(2)
            img_path = os.path.join(img_dir, img_file)
            # Buscar caption en la siguiente línea (bold text)
            caption = ""
            if i + 2 < len(lines):
                next_line = lines[i + 1].strip()
                if not next_line:
                    next_line = lines[i + 2].strip() if i + 2 < len(lines) else ""
                cap_match = re.match(r"\*\*(.+?)\*\*\s*(.*)", next_line)
                if cap_match:
                    caption = cap_match.group(1)
                    rest = cap_match.group(2)
                    if rest:
                        caption += " " + rest
                    # Skip the caption line(s)
                    i += 1
                    while i < len(lines) and not lines[i].strip():
                        i += 1
                    if i < len(lines) and re.match(r"\*\*", lines[i].strip()):
                        i += 1

            add_image_to_doc(doc, img_path, caption)
            i += 1
            continue

        # Headings
        if stripped.startswith("#"):
            level = len(stripped.split(" ")[0])
            text = stripped.lstrip("#").strip()
            adj_level = min(level + heading_offset, 4)
            if adj_level < 1:
                adj_level = 1
            doc.add_heading(text, level=adj_level)
            i += 1
            continue

        # Table (collect consecutive | lines)
        if stripped.startswith("|"):
            # Check if previous line was a bold caption
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = parse_md_table(table_lines)
            add_table_to_doc(doc, rows)
            continue

        # Bold-only line (table/figure caption)
        bold_match = re.match(r"^\*\*(.+?)\*\*\.?\s*(.*)", stripped)
        if bold_match and not stripped.startswith("- "):
            caption_label = bold_match.group(1)
            caption_rest = bold_match.group(2)
            full_caption = caption_label
            if caption_rest:
                full_caption += " " + caption_rest
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(full_caption)
            run.font.size = Pt(9)
            run.italic = True
            i += 1
            continue

        # Bullet list
        if stripped.startswith("- "):
            text = stripped[2:]
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            text = re.sub(r"`(.+?)`", r"\1", text)
            doc.add_paragraph(text, style="List Bullet")
            i += 1
            continue

        # Numbered list
        num_match = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if num_match:
            text = num_match.group(2)
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            text = re.sub(r"`(.+?)`", r"\1", text)
            doc.add_paragraph(text, style="List Number")
            i += 1
            continue

        # Horizontal rule
        if stripped == "---":
            i += 1
            continue

        # Regular paragraph
        text = stripped
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        text = re.sub(r"`(.+?)`", r"\1", text)
        doc.add_paragraph(text)
        i += 1


def read_md(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Genera el informe DOCX del laboratorio.")
    parser.add_argument("--output", default="Informe_Tarea2_Aprendizaje_No_Supervisado.docx")
    args = parser.parse_args()

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Portada ──
    add_cover(doc)

    # ── 1. Introducción ──
    doc.add_heading("1. Introducción", level=1)

    doc.add_heading("Contexto general del aprendizaje no supervisado", level=2)
    doc.add_paragraph(
        "El aprendizaje no supervisado comprende un conjunto de técnicas de "
        "análisis de datos que buscan descubrir patrones, estructuras o "
        "relaciones en datos sin etiquetas predefinidas. A diferencia del "
        "aprendizaje supervisado, no se dispone de una variable objetivo; el "
        "algoritmo debe identificar por sí mismo la organización inherente "
        "de los datos."
    )
    doc.add_paragraph(
        "Entre las tareas principales del aprendizaje no supervisado se "
        "encuentran:"
    )
    doc.add_paragraph(
        "Reducción de dimensionalidad: representar datos de alta dimensión "
        "en espacios de menor dimensión preservando la información más "
        "relevante (SVD, PCA, ICA).", style="List Bullet")
    doc.add_paragraph(
        "Visualización: proyectar datos multidimensionales a 2D o 3D para "
        "exploración visual (t-SNE, UMAP).", style="List Bullet")
    doc.add_paragraph(
        "Separación de fuentes: descomponer señales observadas en "
        "componentes independientes subyacentes (ICA).", style="List Bullet")
    doc.add_paragraph(
        "Estas técnicas son fundamentales en el análisis moderno de datos, "
        "ya que permiten explorar, comprender y preprocesar conjuntos de "
        "datos complejos antes de aplicar modelos predictivos o de toma de "
        "decisiones."
    )

    doc.add_heading("Objetivos del trabajo", level=2)
    doc.add_paragraph(
        "Comprender el funcionamiento teórico de cuatro algoritmos de "
        "aprendizaje no supervisado: SVD, t-SNE, UMAP e ICA.",
        style="List Number")
    doc.add_paragraph(
        "Identificar los principales usos, aplicaciones y limitaciones de "
        "cada algoritmo.", style="List Number")
    doc.add_paragraph(
        "Aplicar cada algoritmo a un conjunto de datos real, analizando e "
        "interpretando los resultados obtenidos.", style="List Number")
    doc.add_paragraph(
        "Comparar los algoritmos entre sí, destacando ventajas, "
        "limitaciones y contextos de uso apropiados.", style="List Number")

    # ── 2. Desarrollo ──
    doc.add_heading("2. Desarrollo", level=1)

    # 2.1 SVD
    doc.add_heading("2.1 SVD (Descomposición en Valores Singulares)", level=2)
    svd_md = read_md(os.path.join(BASE_DIR, "outputs", "svd", "resumen_svd.md"))
    svd_md = re.sub(r"^#\s+.*\n", "", svd_md)
    process_md_content(doc, svd_md,
                       os.path.join(BASE_DIR, "outputs", "svd"),
                       heading_offset=1)

    doc.add_page_break()

    # 2.2 t-SNE
    doc.add_heading("2.2 t-SNE (Incrustación Estocástica de Vecinos con "
                    "Distribución t)", level=2)
    tsne_md = read_md(os.path.join(BASE_DIR, "outputs", "tsne", "resumen_tsne.md"))
    tsne_md = re.sub(r"^#\s+.*\n", "", tsne_md)
    process_md_content(doc, tsne_md,
                       os.path.join(BASE_DIR, "outputs", "tsne"),
                       heading_offset=1)

    doc.add_page_break()

    # 2.3 UMAP
    doc.add_heading("2.3 UMAP (Aproximación y Proyección Uniforme de "
                    "Variedades)", level=2)
    umap_md = read_md(os.path.join(BASE_DIR, "outputs", "umap", "resumen_umap.md"))
    umap_md = re.sub(r"^#\s+.*\n", "", umap_md)
    process_md_content(doc, umap_md,
                       os.path.join(BASE_DIR, "outputs", "umap"),
                       heading_offset=1)

    doc.add_page_break()

    # 2.4 ICA
    doc.add_heading("2.4 ICA (Análisis de Componentes Independientes)", level=2)
    ica_md = read_md(os.path.join(BASE_DIR, "outputs", "ica", "resumen_ica.md"))
    ica_md = re.sub(r"^#\s+.*\n", "", ica_md)
    process_md_content(doc, ica_md,
                       os.path.join(BASE_DIR, "outputs", "ica"),
                       heading_offset=1)

    doc.add_page_break()

    # ── 3. Comparación general ──
    doc.add_heading("3. Comparación general", level=1)

    doc.add_heading("Comparación entre algoritmos", level=2)
    comp_rows = [
        ["Aspecto", "SVD", "t-SNE", "UMAP", "ICA"],
        ["Tipo", "Lineal", "No lineal", "No lineal", "Lineal"],
        ["Objetivo", "Reducir dimensionalidad (varianza)",
         "Visualización 2D/3D", "Visualización + reducción",
         "Separar fuentes independientes"],
        ["Preserva", "Varianza global", "Estructura local",
         "Local + global", "Independencia estadística"],
        ["Escalabilidad", "Excelente", "Limitada (O(n²))",
         "Buena (O(n^1.14))", "Buena"],
        ["Datos nuevos", "Sí", "No", "Sí (.transform())",
         "Sí (matriz W)"],
        ["Supuesto clave", "Ninguno especial", "No paramétrico",
         "Datos sobre manifold", "Fuentes no-gaussianas"],
        ["Determinismo", "Pseudoaleatorio", "Estocástico",
         "Estocástico", "Pseudoaleatorio"],
    ]
    add_table_to_doc(doc, comp_rows,
                     "Tabla 1. Comparación de características entre los "
                     "cuatro algoritmos estudiados.")

    doc.add_heading("Ventajas y limitaciones", level=2)
    vl_rows = [
        ["Algoritmo", "Ventajas", "Limitaciones"],
        ["SVD", "Eficiente con matrices dispersas; interpretable; "
         "base matemática sólida",
         "No captura relaciones no lineales; sensible a valores extremos"],
        ["t-SNE", "Excelente visualización de agrupamientos; revela "
         "estructura local fina",
         "Lento para datasets grandes; no preserva distancias globales; "
         "no transforma datos nuevos"],
        ["UMAP", "Rápido; preserva estructura global y local; soporta "
         "datos nuevos",
         "Depende de hiperparámetros; fundamento teórico complejo"],
        ["ICA", "Separa fuentes independientes; útil para señales "
         "biomédicas; interpretable",
         "Requiere tantos sensores como fuentes; asume mezcla lineal"],
    ]
    add_table_to_doc(doc, vl_rows,
                     "Tabla 2. Ventajas y limitaciones de cada algoritmo.")

    # ── 4. Conclusiones ──
    doc.add_heading("4. Conclusiones", level=1)

    doc.add_heading("Principales aprendizajes", level=2)
    doc.add_paragraph(
        "Cada algoritmo tiene un propósito distinto dentro del aprendizaje "
        "no supervisado: SVD para factorización y reducción, t-SNE y UMAP "
        "para visualización, e ICA para separación de fuentes.",
        style="List Bullet")
    doc.add_paragraph(
        "La elección del algoritmo depende del objetivo del análisis: si se "
        "busca comprimir información (SVD), explorar visualmente "
        "(t-SNE/UMAP) o descomponer señales mixtas (ICA).",
        style="List Bullet")
    doc.add_paragraph(
        "Los hiperparámetros (perplejidad en t-SNE, n_neighbors en UMAP, "
        "número de componentes en SVD) tienen un impacto significativo en "
        "los resultados y deben explorarse sistemáticamente.",
        style="List Bullet")

    doc.add_heading("Dificultades encontradas", level=2)
    doc.add_paragraph(
        "La descarga y lectura de datos en formato WFDB (PhysioNet) "
        "requiere familiarizarse con la librería wfdb y el formato de "
        "anotaciones.", style="List Bullet")
    doc.add_paragraph(
        "t-SNE es computacionalmente costoso y sensible a la perplejidad; "
        "requiere experimentar con varios valores para obtener "
        "visualizaciones informativas.", style="List Bullet")
    doc.add_paragraph(
        "La interpretación de los componentes ICA en señales ECG requiere "
        "conocimiento del dominio (cardiología) para validar si la "
        "separación de fuentes tiene sentido fisiológico.",
        style="List Bullet")

    doc.add_heading("Reflexión final del grupo", level=2)
    doc.add_paragraph(
        "Los algoritmos de aprendizaje no supervisado son herramientas "
        "complementarias, no competidoras. En un flujo de trabajo real de "
        "análisis de datos, es común usar SVD para preprocesar y comprimir, "
        "t-SNE o UMAP para explorar visualmente los patrones descubiertos, "
        "e ICA cuando se necesita separar señales mezcladas. La comprensión "
        "de sus fundamentos teóricos, supuestos y limitaciones es esencial "
        "para seleccionar la técnica adecuada y evitar interpretaciones "
        "erróneas de los resultados."
    )

    # ── 5. Referencias ──
    doc.add_heading("5. Referencias", level=1)
    refs = [
        'Golub, G. H., & Van Loan, C. F. (2013). Matrix Computations (4th ed.). Johns Hopkins University Press.',
        'van der Maaten, L., & Hinton, G. (2008). Visualizing Data using t-SNE. Journal of Machine Learning Research, 9, 2579–2605.',
        'McInnes, L., Healy, J., & Melville, J. (2018). UMAP: Uniform Manifold Approximation and Projection for Dimension Reduction. arXiv:1802.03426.',
        'Hyvärinen, A., & Oja, E. (2000). Independent Component Analysis: Algorithms and Applications. Neural Networks, 13(4–5), 411–430.',
        'Pedregosa, F., et al. (2011). Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research, 12, 2825–2830.',
        'GroupLens Research. MovieLens 100K Dataset. https://grouplens.org/datasets/movielens/100k/',
        'UCI Machine Learning Repository. Breast Cancer Wisconsin (Diagnostic) Data Set. https://archive.ics.uci.edu/ml/datasets/Breast+Cancer+Wisconsin+(Diagnostic)',
        'Goldberger, A. L., et al. (2000). PhysioBank, PhysioToolkit, and PhysioNet. Circulation, 101(23), e215–e220. https://physionet.org/',
        'Llamedo, M., & Martínez, J.P. (2018). MIT-BIH Arrhythmia Database P-Wave Annotations. PhysioNet. https://physionet.org/content/pwave/1.0.0/',
        'Documentación scikit-learn: https://scikit-learn.org/stable/',
        'Documentación UMAP: https://umap-learn.readthedocs.io/',
        'Documentación wfdb-python: https://wfdb.readthedocs.io/',
    ]
    for idx, ref in enumerate(refs, 1):
        doc.add_paragraph(f"[{idx}] {ref}")

    # ── Guardar ──
    output_path = os.path.join(BASE_DIR, args.output)
    doc.save(output_path)
    print(f"Informe generado: {output_path}")


if __name__ == "__main__":
    main()
